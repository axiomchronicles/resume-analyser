from docx.enum.text import WD_COLOR_INDEX
from .utils import HighlightSeverity, _build_highlight_rules

from typing import Dict, Sequence, Any
from docx import Document


WORD_COLOR_MAP = {
    HighlightSeverity.WEAK_PHRASE: WD_COLOR_INDEX.RED,       # red background
    HighlightSeverity.ACTION_MISSING: WD_COLOR_INDEX.YELLOW, # yellow
    HighlightSeverity.METRIC_MISSING: WD_COLOR_INDEX.TURQUOISE,  # cyan-ish
}


def _highlight_run_text(run, start: int, end: int, color: WD_COLOR_INDEX):
    original = run.text
    before = original[:start]
    target = original[start:end]
    after = original[end:]

    parent = run._element.getparent()
    run_idx = parent.index(run._element)

    parent.remove(run._element)

    if before:
        new_run = run._r.addprevious(run._r.__class__(run._r))
        new_run.text = before

    new_run = run._r.addprevious(run._r.__class__(run._r))
    new_run.text = target
    new_run.rPr = run._r.rPr
    new_run.rPr.highlight = color

    if after:
        new_run = run._r.addprevious(run._r.__class__(run._r))
        new_run.text = after


def _highlight_phrase_in_paragraph(paragraph, phrase: str, color: WD_COLOR_INDEX):
    if not paragraph.runs:
        return

    text = paragraph.text
    idx = 0

    while True:
        idx = text.lower().find(phrase.lower(), idx)
        if idx == -1:
            break

        remaining = idx
        for run in paragraph.runs:
            run_len = len(run.text)
            if remaining < run_len:
                start_in_run = remaining
                end_in_run = start_in_run + len(phrase)
                _highlight_run_text(run, start_in_run, end_in_run, color)
                break
            remaining -= run_len

        idx += len(phrase)


def _highlight_phrase_in_table(table, phrase: str, color: WD_COLOR_INDEX):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _highlight_phrase_in_paragraph(paragraph, phrase, color)


def highlight_docx(
    input_path: str,
    output_path: str,
    weak_phrases: Sequence[Dict[str, Any]],
    bullets: Sequence[str],
) -> str:
    if not input_path.lower().endswith(".docx"):
        raise ValueError(f"highlight_docx only supports DOCX files, got: {input_path!r}")

    rules = _build_highlight_rules(weak_phrases, bullets)
    if not rules:
        return input_path

    doc = Document(input_path)

    for rule in rules:
        phrase = rule.phrase
        if not phrase:
            continue

        color = WORD_COLOR_MAP.get(rule.severity, WD_COLOR_INDEX.RED)

        # paragraphs
        for paragraph in doc.paragraphs:
            _highlight_phrase_in_paragraph(paragraph, phrase, color)

        # tables
        for table in doc.tables:
            _highlight_phrase_in_table(table, phrase, color)

    doc.save(output_path)
    return output_path
